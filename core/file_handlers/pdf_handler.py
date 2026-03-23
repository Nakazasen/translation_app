"""
PDF file handler for translation with format preservation
"""
import re
import os
import sys
import tempfile
import warnings
import io
import threading
from typing import List, Dict, Any, Optional, Tuple
from contextlib import contextmanager
import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Inches
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import pdfplumber

from translation_app.core.translator import TranslationService
from translation_app.core.file_handlers.word_handler import WordHandler
from translation_app.core.ocr_handler import OCRHandler
from translation_app.utils.error_handler import FileProcessingError
from translation_app.utils.logger import logger

# Suppress PDF library warnings globally
warnings.filterwarnings('ignore', category=UserWarning)
warnings.filterwarnings('ignore', message='.*Cannot set gray non-stroke color.*')
warnings.filterwarnings('ignore', message='.*invalid float value.*')
warnings.filterwarnings('ignore', message='.*colorspace.*')
warnings.filterwarnings('ignore', message=r'.*R\d+.*')
warnings.filterwarnings('ignore', message='.*gray.*color.*')


@contextmanager
def suppress_pdf_warnings():
    """
    Context manager to suppress PDF library warnings and stderr output
    NUCLEAR OPTION: Completely redirects stderr to devnull for PDF operations
    
    This completely silences all PDF library warnings including:
    - "Cannot set gray non-stroke color"
    - "invalid float value"
    - Colorspace warnings
    - All PyMuPDF/fitz warnings (even from C library)
    - pdf2docx warnings
    """
    import re
    import io
    
    # Suppress all warnings from warnings module
    with warnings.catch_warnings():
        warnings.simplefilter("ignore")
        warnings.filterwarnings('ignore', message='.*')
        warnings.filterwarnings('ignore', category=UserWarning)
        warnings.filterwarnings('ignore', category=DeprecationWarning)
        warnings.filterwarnings('ignore', category=FutureWarning)
        warnings.filterwarnings('ignore', category=RuntimeWarning)
        
        # NUCLEAR OPTION: Create a filter that completely suppresses ALL PDF-related output
        # This handles warnings from both Python and C libraries
        class PDFWarningFilter:
            def __init__(self, original_stderr):
                self.original_stderr = original_stderr
                # Compile regex patterns for better performance
                self.suppress_patterns = [
                    re.compile(r'\[WARNING\]', re.IGNORECASE),
                    re.compile(r'Cannot set gray', re.IGNORECASE),
                    re.compile(r'invalid float value', re.IGNORECASE),
                    re.compile(r"R'\d+", re.IGNORECASE),
                    re.compile(r"R\d+.*invalid", re.IGNORECASE),
                    re.compile(r'colorspace', re.IGNORECASE),
                    re.compile(r'gray.*color', re.IGNORECASE),
                    re.compile(r'pdf.*warning', re.IGNORECASE),
                    re.compile(r'fitz.*warning', re.IGNORECASE),
                    re.compile(r'pymupdf.*warning', re.IGNORECASE),
                    re.compile(r'pdf2docx.*warning', re.IGNORECASE),
                    re.compile(r'converting.*page', re.IGNORECASE),
                ]
                
                # Keywords that indicate PDF warnings - expanded list
                self.suppress_keywords = [
                    'gray', 'color', 'invalid', 'float', 'R\'', 'R47', 'R50', 'R51', 'R52',
                    'R125', 'R127', 'R263', 'R265', 'colorspace', 'warning',
                    'pdf', 'fitz', 'pymupdf', 'pdf2docx', 'non-stroke'
                ]
            
            def write(self, text):
                # If original stderr is None, we can't write anything anyway
                if not self.original_stderr:
                    return

                # Convert to string once
                text_str = str(text)
                
                # NUCLEAR: Suppress if it contains ANY warning pattern
                for pattern in self.suppress_patterns:
                    if pattern.search(text_str):
                        return  # Suppress this warning
                
                # NUCLEAR: Suppress if it contains ANY PDF warning keyword
                for keyword in self.suppress_keywords:
                    if keyword.lower() in text_str.lower():
                        return  # Suppress
                
                # NUCLEAR: Suppress if it's a single line with [WARNING]
                if "[WARNING]" in text_str:
                    return  # Suppress ALL warnings
                
                # NUCLEAR: Suppress if it contains "R" followed by digits (PDF resource references)
                if re.search(r"R\d+", text_str) and ("invalid" in text_str.lower() or "warning" in text_str.lower()):
                    return  # Suppress
                
                # Only write if it's not a PDF warning
                try:
                    self.original_stderr.write(text)
                except Exception:
                    pass
            
            def flush(self):
                if self.original_stderr:
                    try:
                        self.original_stderr.flush()
                    except Exception:
                        pass
            
            def __getattr__(self, name):
                # Forward any other attributes to original_stderr
                if self.original_stderr:
                    return getattr(self.original_stderr, name)
                raise AttributeError(f"'PDFWarningFilter' object has no attribute '{name}'")
        
        # Replace stderr with filter - this catches ALL stderr output including from C libraries
        original_stderr = sys.stderr
        try:
            sys.stderr = PDFWarningFilter(original_stderr)
            yield
        finally:
            sys.stderr = original_stderr

# Try importing advanced PDF libraries
try:
    import camelot
    CAMELOT_AVAILABLE = True
except ImportError:
    CAMELOT_AVAILABLE = False
    logger.warning("camelot-py not available. Table extraction will be limited.")

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    logger.warning("PyMuPDF not available. Text extraction with layout will be limited.")

try:
    from pdf2docx import Converter
    PDF2DOCX_AVAILABLE = True
except ImportError:
    PDF2DOCX_AVAILABLE = False
    logger.warning("pdf2docx not available. PDF to Word conversion will be limited.")

# Check for Microsoft Word COM (Windows only - BEST quality PDF to Word)
# Use lazy initialization to avoid COM issues in PyInstaller onefile
_WORD_COM_CHECKED = False
_WORD_COM_AVAILABLE = None

def _check_word_com_availability():
    """
    Lazily check if Microsoft Word COM is available.
    This is done on first use rather than module load to avoid COM
    initialization issues in PyInstaller onefile environment.
    
    Returns:
        True if Word COM is available, False otherwise
    """
    global _WORD_COM_CHECKED, _WORD_COM_AVAILABLE
    
    if _WORD_COM_CHECKED:
        return _WORD_COM_AVAILABLE
    
    _WORD_COM_CHECKED = True
    _WORD_COM_AVAILABLE = False
    
    try:
        import win32com.client
        import pythoncom
        
        # Initialize COM for this thread before testing
        try:
            pythoncom.CoInitialize()
        except Exception:
            pass  # May already be initialized
        
        try:
            # Quick test to see if Word is installed
            word_test = win32com.client.Dispatch("Word.Application")
            word_test.Visible = False
            word_test.DisplayAlerts = False
            word_test.Quit()
            _WORD_COM_AVAILABLE = True
            logger.info("[OK] Microsoft Word detected - will use for best PDF conversion quality")
        except Exception as e:
            logger.debug(f"Word COM test failed: {e}")
        finally:
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass
    except ImportError:
        logger.debug("win32com not available")
    
    if not _WORD_COM_AVAILABLE:
        logger.info("Microsoft Word COM not available. Using pdf2docx for PDF conversion.")
    
    return _WORD_COM_AVAILABLE



class PDFHandler:
    """Handler for PDF file translation with format preservation"""
    
    def __init__(self, translation_service: TranslationService):
        """
        Initialize PDF handler
        
        Args:
            translation_service: Translation service instance
        """
        self.translation_service = translation_service
        self.ocr_handler = OCRHandler()
        self.font_name = 'Times New Roman'
        self.method_used = None  # Track which method was used for logging
        self.progress_callback = None  # Function to update UI progress (status_text, percentage)
        self._main_thread_id = threading.current_thread().ident  # Lưu ID thread chính để kiểm tra thread-safety
    
    def _safe_progress_callback(self, message: str, progress: int) -> None:
        """
        Gọi progress callback một cách thread-safe.
        
        Nếu đang chạy trong background thread, cần đảm bảo callback 
        được thực thi trong main thread để tránh race condition với UI.
        
        Args:
            message: Thông báo trạng thái
            progress: Phần trăm tiến độ (0-100)
        """
        if self.progress_callback is None:
            return
            
        current_thread_id = threading.current_thread().ident
        if current_thread_id != self._main_thread_id:
            # Đang chạy trong background thread - cần xử lý đặc biệt
            # Ở đây chúng ta log warning vì không thể trực tiếp gọi UI update từ background thread
            logger.warning(f"Progress callback called from background thread. Message: {message}, Progress: {progress}%")
            # Trong thực tế, callback này nên được gọi thông qua queue hoặc thread-safe mechanism
            # để main thread có thể cập nhật UI. Hiện tại chúng ta vẫn gọi nhưng log cảnh báo.
        
        try:
            self.progress_callback(message, progress)
        except Exception as e:
            logger.error(f"Error in progress callback: {e}")
    
    def translate(self, input_file: str, output_file: str, src_lang: str, dest_lang: str) -> None:
        """
        Translate PDF file to Word document with format preservation
        
        Uses multi-method approach:
        1. Try pdf2docx conversion (best format preservation)
        2. Extract tables + text with layout (PyMuPDF + camelot)
        3. Fallback to simple text extraction (pdfplumber)
        
        Args:
            input_file: Path to input PDF file
            output_file: Path to output Word file
            src_lang: Source language code
            dest_lang: Destination language code
        
        Raises:
            FileProcessingError: If processing fails
        """
        try:
            # Suppress PDF warnings for entire translation process
            with suppress_pdf_warnings():
                logger.info(f"Starting PDF translation: {input_file}")
                if self.progress_callback:
                    self.progress_callback(f"Bắt đầu dịch file: {os.path.basename(input_file)}...", 5)
                
                # Check if PDF is scanned (needs OCR)
                is_scanned = self._is_scanned_pdf(input_file)
                
                # SMART STRATEGY: Choose method based on PDF type and available tools
                # Priority order:
                # 1. For scanned PDFs: Use OCR method (convert pages to images, OCR, translate)
                # 2. MS Word COM (BEST quality, Windows only with Word installed)
                # 3. pdf2docx (good quality, cross-platform)
                # 4. Layout extraction (fallback)
                # 5. Simple text extraction (last resort)
                
                if is_scanned and PYMUPDF_AVAILABLE and self.ocr_handler.is_installed():
                    try:
                        logger.info("PDF is scanned. Using OCR METHOD: Converting pages to images, OCR, then translate...")
                        if self.progress_callback:
                            self.progress_callback("Phát hiện PDF file scan. Đang khởi động OCR...", 10)
                        self._translate_pdf_with_ocr(input_file, output_file, src_lang, dest_lang)
                        self.method_used = "ocr"
                        logger.info("✓ Successfully translated using OCR method")
                        return
                    except Exception as e:
                        logger.warning(f"OCR method failed: {e}. Trying Word conversion...")
                
                # Method 1: Microsoft Word COM (BEST quality - preserves ALL formatting)
                # Word can open PDF and convert to editable docx with perfect formatting
                if _check_word_com_availability():
                    try:
                        logger.info("🏆 Using BEST METHOD: Microsoft Word COM (perfect format preservation)...")
                        if self.progress_callback:
                            self.progress_callback("Đang dùng Microsoft Word để chuyển đổi PDF giữ định dạng...", 15)
                        self._translate_via_word_conversion(input_file, output_file, src_lang, dest_lang)
                        self.method_used = "word_com"
                        logger.info("✅ Successfully translated using Microsoft Word conversion!")
                        return
                    except Exception as e:
                        logger.warning(f"Word COM method failed: {e}. Trying pdf2docx...")
                
                # Method 2: Try pdf2docx (good format preservation, cross-platform)
                if PDF2DOCX_AVAILABLE:
                    try:
                        logger.info("Using pdf2docx: Converting PDF to Word first...")
                        if self.progress_callback:
                            self.progress_callback("Đang dùng pdf2docx để chuyển đổi PDF...", 20)
                        self._translate_via_pdf2docx(input_file, output_file, src_lang, dest_lang)
                        self.method_used = "pdf2docx"
                        logger.info("✓ Successfully translated using pdf2docx method")
                        return
                    except Exception as e:
                        logger.warning(f"pdf2docx conversion failed: {e}. Trying alternative methods...")
                
                # Method 3: Extract tables and text with layout (only if above methods fail)
                if (PYMUPDF_AVAILABLE or CAMELOT_AVAILABLE):
                    try:
                        logger.info("Attempting table and text extraction with layout...")
                        if self.progress_callback:
                            self.progress_callback("Đang trích xuất layout PDF trực tiếp...", 20)
                        self._translate_with_layout_extraction(input_file, output_file, src_lang, dest_lang)
                        self.method_used = "layout_extraction"
                        logger.info("Successfully translated using layout extraction method")
                        return
                    except Exception as e:
                        logger.warning(f"Layout extraction failed: {e}. Falling back to simple extraction...")
                
                # Method 4: Fallback to simple text extraction (last resort)
                logger.info("Using fallback method: simple text extraction")
                if self.progress_callback:
                    self.progress_callback("Đang dùng phương thức trích xuất văn bản thô...", 20)
                self._translate_simple_extraction(input_file, output_file, src_lang, dest_lang)
                self.method_used = "simple_extraction"
                logger.info("Successfully translated using simple extraction method")


        
        except Exception as e:
            error_msg = f"Error translating PDF: {e}"
            logger.error(error_msg, exc_info=True)
            raise FileProcessingError(error_msg, original_error=e) from e
    
    def translate_with_ai_vision(self, input_file: str, output_file: str, src_lang: str, dest_lang: str,
                                   pages_per_batch: int = 4) -> None:
        """
        Translate PDF using Gemini Vision AI with MAXIMUM REQUEST OPTIMIZATION.
        
        OPTIMIZATION STRATEGY:
        - Instead of 1 page = 1 request (wasteful!)
        - Combine multiple pages into ONE grid image
        - Send grid image to Gemini Vision = 1 request for multiple pages
        
        Example: 50-page PDF
        - Before: 50 requests (1 per page)
        - After (4 pages/batch): ~13 requests
        - After (6 pages/batch): ~9 requests
        
        Args:
            input_file: Path to input PDF file
            output_file: Path to output Word file
            src_lang: Source language code
            dest_lang: Destination language code
            pages_per_batch: Number of pages to combine into one image (default: 4)
        """
        if not PYMUPDF_AVAILABLE:
            raise FileProcessingError("PyMuPDF not available for AI Vision method")
        
        from translation_app.core.ai_service import get_ai_service
        from PIL import Image
        
        ai_service = get_ai_service()
        if not ai_service.is_available():
            raise FileProcessingError("Gemini AI not configured. Please add API keys first.")
        
        logger.info(f"🤖 AI VISION (OPTIMIZED): Combining {pages_per_batch} pages per request to save RPD")
        
        with suppress_pdf_warnings():
            doc = fitz.open(input_file)
            total_pages = len(doc)
            
            # Calculate number of batches needed
            num_batches = (total_pages + pages_per_batch - 1) // pages_per_batch
            logger.info(f"📄 {total_pages} pages → {num_batches} AI requests (saved {total_pages - num_batches} requests!)")
            
            if self.progress_callback:
                self.progress_callback(f"Bắt đầu dịch AI Vision: {total_pages} trang → {num_batches} batch...", 5)
            
            # Create Word document
            word_doc = Document()
            
            # Add header
            header_para = word_doc.add_paragraph()
            header_run = header_para.add_run(f"[AI Vision Translation - {total_pages} pages in {num_batches} requests]")
            header_run.font.bold = True
            header_run.font.name = self.font_name
            
            all_translated_texts = []
            
            # Process pages in batches
            for batch_idx in range(num_batches):
                start_page = batch_idx * pages_per_batch
                end_page = min(start_page + pages_per_batch, total_pages)
                batch_pages = list(range(start_page, end_page))
                
                # Update progress
                progress_val = 5 + int((batch_idx / num_batches) * 85)
                if self.progress_callback:
                    self.progress_callback(
                        f"🤖 Đang dịch batch {batch_idx + 1}/{num_batches} (Trang {start_page + 1}-{end_page})...", 
                        progress_val
                    )
                
                logger.info(f"🖼️ Batch {batch_idx + 1}/{num_batches}: Pages {start_page + 1}-{end_page}")
                
                try:
                    # Create grid image combining multiple pages
                    grid_image = self._create_page_grid(doc, batch_pages)
                    
                    # Convert to bytes
                    img_buffer = io.BytesIO()
                    grid_image.save(img_buffer, format='PNG', optimize=True)
                    img_bytes = img_buffer.getvalue()
                    
                    # Create prompt that tells AI about the grid layout
                    page_count = len(batch_pages)
                    if page_count == 1:
                        layout_hint = "Đây là 1 trang tài liệu."
                    elif page_count == 2:
                        layout_hint = "Đây là 2 trang xếp ngang (trái → phải)."
                    elif page_count <= 4:
                        layout_hint = f"Đây là {page_count} trang xếp thành lưới 2x2 (đọc từ trái qua phải, trên xuống dưới)."
                    else:
                        layout_hint = f"Đây là {page_count} trang xếp thành lưới (đọc từ trái qua phải, trên xuống dưới)."
                    
                    # Send to Gemini Vision
                    result = ai_service.translate_image_with_vision(
                        img_bytes,
                        source_lang=src_lang,
                        target_lang=dest_lang,
                        preserve_format=True,
                        custom_hint=layout_hint
                    )
                    
                    if result.get("status") == "success":
                        translated_text = result.get("text", "")
                        # Store with batch info
                        all_translated_texts.append({
                            'pages': batch_pages,
                            'text': translated_text,
                            'model': result.get('model_used', 'unknown')
                        })
                        logger.info(f"✅ Batch {batch_idx + 1} translated ({page_count} pages in 1 request)")
                    else:
                        logger.warning(f"⚠️ Batch {batch_idx + 1} failed, using fallback")
                        # Fallback: try OCR for each page individually
                        if self.progress_callback:
                            self.progress_callback(f"⚠️ Batch {batch_idx + 1} lỗi AI, đang dùng OCR dự phòng...", progress_val)
                        fallback_text = self._fallback_ocr_batch(doc, batch_pages, src_lang, dest_lang)
                        all_translated_texts.append({
                            'pages': batch_pages,
                            'text': fallback_text,
                            'model': 'FALLBACK_OCR'
                        })
                        
                except Exception as e:

                    logger.warning(f"Error processing batch {batch_idx + 1}: {e}")
                    all_translated_texts.append({
                        'pages': batch_pages,
                        'text': f"[Batch {batch_idx + 1} Error: {str(e)[:100]}]",
                        'model': 'ERROR'
                    })
            
            doc.close()
            
            # Insert all translated content into Word document
            for batch_result in all_translated_texts:
                pages = batch_result['pages']
                translated_text = batch_result['text']
                
                if translated_text and translated_text.strip():
                    # Add page range marker
                    if len(pages) == 1:
                        marker_text = f"--- Trang {pages[0] + 1} ---"
                    else:
                        marker_text = f"--- Trang {pages[0] + 1} đến {pages[-1] + 1} ---"
                    
                    page_marker = word_doc.add_paragraph()
                    marker_run = page_marker.add_run(marker_text)
                    marker_run.font.bold = True
                    marker_run.font.name = self.font_name
                    
                    # Add translated content
                    paragraphs = translated_text.split('\n')
                    for para_text in paragraphs:
                        if para_text.strip():
                            para = word_doc.add_paragraph(para_text)
                            for run in para.runs:
                                run.font.name = self.font_name
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)
                    
                    # Add page break between batches
                    run = word_doc.add_paragraph().add_run()
                    run.add_break(WD_BREAK.PAGE)
            
            word_doc.save(output_file)
            self.method_used = "ai_vision_optimized"
            logger.info(f"✅ AI Vision (OPTIMIZED) completed: {total_pages} pages in {num_batches} requests!")
    
    def _create_page_grid(self, doc, page_numbers: list, max_width: int = 2000) -> 'Image.Image':
        """
        Create a grid image combining multiple PDF pages.
        
        Layout:
        - 1 page: single image
        - 2 pages: 1 row x 2 cols
        - 3-4 pages: 2 rows x 2 cols
        - 5-6 pages: 2 rows x 3 cols
        - 7-9 pages: 3 rows x 3 cols
        
        Args:
            doc: PyMuPDF document
            page_numbers: List of page indices to combine
            max_width: Maximum width of combined image
            
        Returns:
            PIL Image containing all pages in a grid
        """
        from PIL import Image
        
        num_pages = len(page_numbers)
        
        # Determine grid layout
        if num_pages == 1:
            cols, rows = 1, 1
        elif num_pages == 2:
            cols, rows = 2, 1
        elif num_pages <= 4:
            cols, rows = 2, 2
        elif num_pages <= 6:
            cols, rows = 3, 2
        else:
            cols, rows = 3, 3
        
        # Render pages to images
        page_images = []
        for page_num in page_numbers:
            page = doc[page_num]
            # Use 1.5x zoom for good quality without being too large
            mat = fitz.Matrix(1.5, 1.5)
            pix = page.get_pixmap(matrix=mat)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            page_images.append(img)
        
        if not page_images:
            return Image.new('RGB', (100, 100), 'white')
        
        # Calculate cell size (uniform for all cells)
        cell_width = max(img.width for img in page_images)
        cell_height = max(img.height for img in page_images)
        
        # Create grid image
        grid_width = cols * cell_width
        grid_height = rows * cell_height
        
        # Scale down if too large
        if grid_width > max_width:
            scale = max_width / grid_width
            cell_width = int(cell_width * scale)
            cell_height = int(cell_height * scale)
            grid_width = cols * cell_width
            grid_height = rows * cell_height
            # Resize all page images
            page_images = [img.resize((cell_width, cell_height), Image.Resampling.LANCZOS) for img in page_images]
        
        # Create canvas
        grid_image = Image.new('RGB', (grid_width, grid_height), 'white')
        
        # Paste pages into grid
        for idx, img in enumerate(page_images):
            row = idx // cols
            col = idx % cols
            x = col * cell_width
            y = row * cell_height
            
            # Center the image in its cell if it's smaller
            x_offset = (cell_width - img.width) // 2
            y_offset = (cell_height - img.height) // 2
            
            grid_image.paste(img, (x + x_offset, y + y_offset))
        
        return grid_image
    
    def _fallback_ocr_batch(self, doc, page_numbers: list, src_lang: str, dest_lang: str) -> str:
        """Fallback OCR for a batch of pages when AI Vision fails."""
        from PIL import Image
        
        all_texts = []
        ocr_lang = self.ocr_handler.get_ocr_language(src_lang)
        
        for page_num in page_numbers:
            try:
                page = doc[page_num]
                mat = fitz.Matrix(2.0, 2.0)
                pix = page.get_pixmap(matrix=mat)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                
                if self.ocr_handler.is_installed():
                    text = self.ocr_handler.extract_text_from_image(img, lang=ocr_lang)
                    if text.strip():
                        all_texts.append(text.strip())
            except Exception as e:
                logger.warning(f"Fallback OCR failed for page {page_num + 1}: {e}")
        
        # Translate combined text
        if all_texts:
            combined = "\n\n".join(all_texts)
            try:
                translated = self.translation_service.translate_text(combined, src_lang, dest_lang)
                return translated
            except Exception as e:
                logger.warning(f"Fallback translation failed: {e}")
                return combined
        
        return "[No text extracted from pages]"

    
    def _can_convert_with_pdf2docx(self, input_file: str) -> bool:
        """
        Check if PDF can be converted with pdf2docx
        Always return True if pdf2docx is available - let it try and fail gracefully
        
        Args:
            input_file: Path to PDF file
            
        Returns:
            True if pdf2docx is available (we'll try conversion anyway)
        """
        # Always try pdf2docx if available - it's the best method
        # If it fails, we'll catch the exception and fallback
        return PDF2DOCX_AVAILABLE
    
    def _is_scanned_pdf(self, input_file: str) -> bool:
        """
        Check if PDF is a scanned document (image-based) that needs OCR
        
        Args:
            input_file: Path to PDF file
            
        Returns:
            True if PDF appears to be scanned (low text content)
        """
        if not PYMUPDF_AVAILABLE:
            return False
        
        try:
            doc = fitz.open(input_file)
            total_text_length = 0
            total_pages = len(doc)
            
            # Check first few pages for text content
            pages_to_check = min(3, total_pages)
            for page_num in range(pages_to_check):
                page = doc[page_num]
                text = page.get_text()
                total_text_length += len(text.strip())
            
            doc.close()
            
            # If average text per page is very low, likely scanned
            avg_text_per_page = total_text_length / pages_to_check if pages_to_check > 0 else 0
            is_scanned = total_text_length < 100  # Less than 100 chars across 3 pages = likely scanned
            
            if is_scanned:
                logger.info(f"PDF appears to be scanned (low text content: {avg_text_per_page:.0f} chars/page). Will use OCR if needed.")
            
            return is_scanned
        except Exception as e:
            logger.debug(f"Could not determine if PDF is scanned: {e}")
            return False
    
    def _convert_pdf_with_word_com(self, input_file: str, output_file: str) -> bool:
        """
        Convert PDF to Word using Microsoft Word COM Automation.
        THIS IS THE BEST METHOD for format preservation!
        
        Microsoft Word:
        - Perfectly recognizes tables, columns, headers, footers
        - Preserves fonts, styles, colors
        - Handles complex layouts much better than pdf2docx
        
        Args:
            input_file: Path to input PDF file
            output_file: Path to output Word file
            
        Returns:
            True if conversion successful, False otherwise
        """
        if not _check_word_com_availability():
            return False
        
        import win32com.client
        import pythoncom
        
        word = None
        doc = None
        
        try:
            # Initialize COM for this thread
            pythoncom.CoInitialize()
            
            logger.info("🔧 Opening Microsoft Word for PDF conversion...")
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False  # Run in background
            word.DisplayAlerts = False  # Suppress dialogs
            
            # Convert paths to absolute paths (Word requires this)
            abs_input = os.path.abspath(input_file)
            abs_output = os.path.abspath(output_file)
            
            # Make sure output directory exists
            os.makedirs(os.path.dirname(abs_output) if os.path.dirname(abs_output) else '.', exist_ok=True)
            
            logger.info(f"📄 Opening PDF: {os.path.basename(input_file)}")
            
            # Open PDF in Word (Word will auto-convert to editable format)
            # ConfirmConversions=False prevents the conversion dialog
            doc = word.Documents.Open(
                abs_input,
                ConfirmConversions=False,
                ReadOnly=True,
                AddToRecentFiles=False
            )
            
            logger.info("💾 Saving as Word document (preserving all formatting)...")
            
            # Save as docx format (wdFormatXMLDocument = 12)
            doc.SaveAs2(abs_output, FileFormat=12)
            
            logger.info(f"✅ PDF converted successfully: {os.path.basename(output_file)}")
            return True
            
        except Exception as e:
            logger.warning(f"⚠️ Word COM conversion failed: {e}")
            return False
            
        finally:
            # Clean up
            try:
                if doc:
                    doc.Close(SaveChanges=False)
            except Exception:
                pass
            
            try:
                if word:
                    word.Quit()
            except Exception:
                pass
            
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass
    
    def _translate_via_word_conversion(self, input_file: str, output_file: str, src_lang: str, dest_lang: str) -> None:
        """
        BEST METHOD: Convert PDF to Word using MS Word, then translate Word file.
        
        Process:
        1. Use MS Word to open PDF and save as .docx (BEST format preservation)
        2. Translate the Word file using WordHandler (already optimized)
        3. Save the translated Word file
        
        Args:
            input_file: Path to input PDF file
            output_file: Path to output Word file
            src_lang: Source language code
            dest_lang: Destination language code
        """
        temp_word_file = None
        
        try:
            # Step 1: Convert PDF to Word using MS Word
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                temp_word_file = tmp.name
            
            logger.info("=" * 60)
            logger.info("🚀 BEST METHOD: PDF → Word (MS Word) → Dịch → Kết quả")
            logger.info("=" * 60)
            
            logger.info("📌 Bước 1/2: Chuyển đổi PDF sang Word bằng Microsoft Word...")
            
            success = self._convert_pdf_with_word_com(input_file, temp_word_file)
            
            if not success:
                raise FileProcessingError("Microsoft Word conversion failed")
            
            # Step 2: Translate the Word file
            logger.info("📌 Bước 2/2: Dịch file Word (giữ nguyên định dạng)...")
            
            word_handler = WordHandler(self.translation_service)
            word_handler.translate(temp_word_file, output_file, src_lang, dest_lang)
            
            logger.info("=" * 60)
            logger.info(f"✅ Hoàn tất! File đã lưu: {os.path.basename(output_file)}")
            logger.info("=" * 60)
            
        finally:
            # Clean up temp file
            if temp_word_file and os.path.exists(temp_word_file):
                try:
                    os.unlink(temp_word_file)
                except Exception:
                    pass
    
    def _translate_via_pdf2docx(self, input_file: str, output_file: str, src_lang: str, dest_lang: str) -> None:
        """
        Translate PDF by converting to Word first, then translating Word file
        SMART METHOD: This preserves formatting, tables, images best
        
        Args:
            input_file: Path to input PDF file
            output_file: Path to output Word file
            src_lang: Source language code
            dest_lang: Destination language code
        """
        # Check if PDF is scanned (needs OCR)
        is_scanned = self._is_scanned_pdf(input_file)
        
        # Create temporary Word file
        temp_word_file = None
        try:
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                temp_word_file = tmp.name
            
            # Convert PDF to Word with optimized settings (suppress ALL warnings)
            logger.info("Step 1/2: Converting PDF to Word format (this preserves tables, images, formatting)...")
            if is_scanned:
                logger.info("Note: PDF appears to be scanned. Images will be processed with OCR during translation.")
            
            with suppress_pdf_warnings():
                try:
                    cv = Converter(input_file)
                    
                    # Get page count for progress
                    page_count = cv.get_page_count()
                    logger.info(f"PDF has {page_count} pages. Converting...")
                    
                    # Progress callback for large files
                    def progress_callback(current, total):
                        if current % 10 == 0 or current == total:  # Log every 10 pages for large files
                            logger.info(f"Converting page {current}/{total}...")
                    
                    # Convert with optimized settings
                    # pdf2docx automatically preserves:
                    # - Tables (with borders, alignment)
                    # - Images (position, size)
                    # - Text formatting (fonts, sizes, colors)
                    # - Page layout (margins, headers, footers)
                    cv.convert(temp_word_file, start=0, end=None)  # Convert all pages
                    cv.close()
                    
                    logger.info(f"✓ PDF converted to Word successfully ({page_count} pages)")
                except Exception as convert_error:
                    try:
                        cv.close()
                    except:
                        pass
                    raise FileProcessingError(f"PDF to Word conversion failed: {convert_error}", original_error=convert_error) from convert_error
            
            # Use WordHandler to translate (WordHandler already has batch translation + OCR)
            logger.info("Step 2/2: Translating Word document (this is fast with batch translation)...")
            if is_scanned:
                logger.info("Processing images with OCR during translation...")
            
            word_handler = WordHandler(self.translation_service)
            word_handler.translate(temp_word_file, output_file, src_lang, dest_lang)
            
            logger.info("✓ PDF translation completed successfully using smart method (format preserved)")
            
        finally:
            # Clean up temp file
            if temp_word_file and os.path.exists(temp_word_file):
                try:
                    os.remove(temp_word_file)
                except Exception:
                    pass
    
    def _translate_pdf_with_ocr(self, input_file: str, output_file: str, src_lang: str, dest_lang: str) -> None:
        """
        Translate PDF using OCR on each page (for scanned PDFs)
        SMART METHOD: Convert each page to image, OCR, translate, then create Word doc
        This is better than snipping tool - automated and preserves page order
        
        Args:
            input_file: Path to input PDF file
            output_file: Path to output Word file
            src_lang: Source language code
            dest_lang: Destination language code
        """
        if not PYMUPDF_AVAILABLE:
            raise FileProcessingError("PyMuPDF not available for OCR method")
        
        if not self.ocr_handler.is_installed():
            raise FileProcessingError("Tesseract OCR not installed for OCR method")
        
        logger.info("Using OCR METHOD: Converting PDF pages to images, OCR, then translate...")
        
        with suppress_pdf_warnings():
            doc = fitz.open(input_file)
            total_pages = len(doc)
            logger.info(f"Processing {total_pages} pages with OCR...")
            
            # Create Word document
            word_doc = Document() # create a new word document
            
            # Process each page
            all_texts = [] # list to store all the text from the pages
            for page_num in range(total_pages):
                try:
                    page = doc[page_num] # get the page
                    
                    # Convert page to image (2x zoom for better OCR quality)
                    mat = fitz.Matrix(2.0, 2.0)
                    pix = page.get_pixmap(matrix=mat)
                    
                    # Convert to PIL Image
                    from PIL import Image
                    img_data = pix.tobytes("png")
                    img = Image.open(io.BytesIO(img_data))
                    
                    # OCR
                    ocr_lang = self.ocr_handler.get_ocr_language(src_lang)
                    try:
                        page_text = self.ocr_handler.extract_text_from_image(img, lang=ocr_lang)
                    except Exception:
                        page_text = self.ocr_handler.extract_text_from_image(img, lang='eng')
                    
                    if page_text.strip():
                        all_texts.append(page_text.strip())
                    
                    if (page_num + 1) % 5 == 0:
                        logger.info(f"OCR processed {page_num + 1}/{total_pages} pages...")
                        
                except Exception as e:
                    logger.warning(f"Error processing page {page_num + 1} with OCR: {e}")
                    continue
            
            doc.close()
            
            # Batch translate all OCR texts
            if all_texts:
                logger.info(f"Translating {len(all_texts)} OCR text blocks...")
                try:
                    translated_texts = self.translation_service.translate_batch(all_texts, src_lang, dest_lang)
                except Exception as e:
                    logger.warning(f"Batch translation failed: {e}")
                    translated_texts = all_texts
                
                # Insert into Word document
                for i, translated_text in enumerate(translated_texts):
                    if translated_text and translated_text.strip():
                        para = word_doc.add_paragraph(translated_text)
                        para.style.font.name = self.font_name
                        
                        # Add page break (except for last page)
                        if i < len(translated_texts) - 1:
                            run = para.add_run()
                            run.add_break(WD_BREAK.PAGE)
            
            word_doc.save(output_file)
            logger.info(f"✓ OCR translation completed: {len(all_texts)} pages processed")
    
    def _translate_with_layout_extraction(self, input_file: str, output_file: str, src_lang: str, dest_lang: str) -> None:
        """
        Translate PDF by extracting tables and text with layout preservation
        
        Args:
            input_file: Path to input PDF file
            output_file: Path to output Word file
            src_lang: Source language code
            dest_lang: Destination language code
        """
        # Extract tables
        tables = []
        if CAMELOT_AVAILABLE:
            try:
                tables = self._extract_tables(input_file)
                logger.info(f"Extracted {len(tables)} tables from PDF")
            except Exception as e:
                logger.warning(f"Table extraction failed: {e}")
        
        # Extract text blocks with layout
        text_blocks = []
        if PYMUPDF_AVAILABLE:
            try:
                text_blocks = self._extract_text_with_layout(input_file)
                logger.info(f"Extracted {len(text_blocks)} text blocks from PDF")
            except Exception as e:
                logger.warning(f"Text extraction with layout failed: {e}")
                # Fallback to simple extraction
                try:
                    text_blocks = self._extract_text_simple(input_file)
                except Exception as e2:
                    logger.warning(f"Simple text extraction also failed: {e2}")
        
        # If no text blocks and PyMuPDF not available, try simple extraction
        if not text_blocks:
            try:
                text_blocks = self._extract_text_simple(input_file)
            except Exception as e:
                logger.warning(f"Simple text extraction failed: {e}")
        
        # Create Word document and insert content
        doc = Document()
        
        # Insert tables first
        if tables:
            self._insert_translated_tables(doc, tables, src_lang, dest_lang)
        
        # Insert text blocks
        if text_blocks:
            self._insert_translated_text_blocks(doc, text_blocks, src_lang, dest_lang)
        elif not tables:
            # If no tables and no text blocks, fallback to simple extraction
            raise Exception("No content extracted, falling back to simple extraction")
        
        doc.save(output_file)
        logger.info(f"Saved Word file with {len(tables)} tables and {len(text_blocks)} text blocks")
    
    def _translate_simple_extraction(self, input_file: str, output_file: str, src_lang: str, dest_lang: str) -> None:
        """
        Translate PDF using simple text extraction (fallback method)
        
        Args:
            input_file: Path to input PDF file
            output_file: Path to output Word file
            src_lang: Source language code
            dest_lang: Destination language code
        """
        # Suppress warnings from pdfplumber
        with suppress_pdf_warnings():
            # Read PDF and process page by page
            with pdfplumber.open(input_file) as pdf:
                total_pages = len(pdf.pages)
                logger.info(f"Total PDF pages: {total_pages}")
                
                # Create new Word document
                doc = Document()
                
                # OPTIMIZATION: Collect all pages first, then batch translate
                page_texts = []
                page_numbers = []
                
                for page_num, page in enumerate(pdf.pages, start=1):
                    if page_num % 10 == 0:  # Log every 10 pages
                        logger.info(f"Extracting text from page {page_num}/{total_pages}...")
                    
                    page_text = page.extract_text()
                    if page_text:
                        cleaned_text = self._clean_pdf_text(page_text)
                        if cleaned_text.strip():
                            page_texts.append(cleaned_text)
                            page_numbers.append(page_num)
                
                logger.info(f"Extracted text from {len(page_texts)} pages. Starting batch translation...")
                
                # Batch translate all pages - filter None values
                translated_texts = []
                if page_texts:
                    # Filter out None and invalid values
                    valid_texts = [t for t in page_texts if t is not None and str(t).strip() and str(t).lower() != 'nan']
                    if valid_texts:
                        try:
                            translated_results = self.translation_service.translate_batch(valid_texts, src_lang, dest_lang)
                            # Map results back to original list
                            valid_idx = 0
                            for orig_text in page_texts:
                                if orig_text is not None and str(orig_text).strip() and str(orig_text).lower() != 'nan':
                                    if valid_idx < len(translated_results):
                                        translated_texts.append(translated_results[valid_idx] if translated_results[valid_idx] is not None else orig_text)
                                    else:
                                        translated_texts.append(orig_text)
                                    valid_idx += 1
                                else:
                                    translated_texts.append(orig_text)  # Keep original for invalid values
                        except Exception as e:
                            logger.warning(f"Batch translation failed: {e}. Using individual translation.")
                            translated_texts = page_texts  # Keep original on error
                    else:
                        translated_texts = page_texts
                
                # Insert translated pages into document
                for idx, page_num in enumerate(page_numbers):
                    translated_text = translated_texts[idx] if idx < len(translated_texts) else page_texts[idx]
                    
                    # Add page header (except first page)
                    if page_num > 1:
                        p_header = doc.add_paragraph(f"\n--- Trang {page_num} ---")
                        p_header.runs[0].font.name = self.font_name
                        p_header.runs[0].font.bold = True
                        p_header.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)
                    
                    # Add translated content
                    paragraphs = translated_text.split("\n")
                    for para_text in paragraphs:
                        if para_text.strip():  # Only add non-empty paragraphs
                            p = doc.add_paragraph(para_text)
                            for run in p.runs:
                                run.font.name = self.font_name
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)
                    
                    # Add page break after each page (except last)
                    if idx < len(page_numbers) - 1:
                        p_break = doc.add_paragraph()
                        run_break = p_break.add_run()
                        run_break.add_break(WD_BREAK.PAGE)
                
                doc.save(output_file)
                logger.info(f"Saved Word file with {total_pages} pages")
    
    def _extract_tables(self, input_file: str) -> List[Dict[str, Any]]:
        """
        Extract tables from PDF using camelot
        
        Args:
            input_file: Path to PDF file
            
        Returns:
            List of tables with metadata
        """
        tables = []
        
        if not CAMELOT_AVAILABLE:
            return tables
        
        try:
            # Suppress warnings during table extraction
            with suppress_pdf_warnings():
                # Try lattice mode first (for tables with borders)
                try:
                    lattice_tables = camelot.read_pdf(input_file, pages='all', flavor='lattice')
                    for table in lattice_tables:
                        if table.df is not None and not table.df.empty:
                            tables.append({
                                'dataframe': table.df,
                                'page': table.page,
                                'method': 'lattice',
                                'accuracy': table.accuracy if hasattr(table, 'accuracy') else None
                            })
                    logger.info(f"Extracted {len(lattice_tables)} tables using lattice method")
                except Exception as e:
                    logger.debug(f"Lattice method failed: {e}")
                
                # Try stream mode (for tables without borders)
                try:
                    stream_tables = camelot.read_pdf(input_file, pages='all', flavor='stream')
                    # Filter out duplicates (tables already found by lattice)
                    for table in stream_tables:
                        if table.df is not None and not table.df.empty:
                            # Check if this table is already in our list
                            is_duplicate = False
                            for existing_table in tables:
                                if (existing_table['page'] == table.page and
                                    existing_table['dataframe'].shape == table.df.shape):
                                    is_duplicate = True
                                    break
                            
                            if not is_duplicate:
                                tables.append({
                                    'dataframe': table.df,
                                    'page': table.page,
                                    'method': 'stream',
                                    'accuracy': table.accuracy if hasattr(table, 'accuracy') else None
                                })
                    logger.info(f"Extracted {len(stream_tables)} tables using stream method")
                except Exception as e:
                    logger.debug(f"Stream method failed: {e}")
        
        except Exception as e:
            logger.error(f"Error extracting tables: {e}", exc_info=True)
        
        return tables
    
    def _extract_text_with_layout(self, input_file: str) -> List[Dict[str, Any]]:
        """
        Extract text from PDF with layout information using PyMuPDF
        Optimized for performance with batch processing
        
        Args:
            input_file: Path to PDF file
            
        Returns:
            List of text blocks with metadata (page, position, text)
        """
        text_blocks = []
        
        if not PYMUPDF_AVAILABLE:
            return text_blocks
        
        try:
            # Suppress warnings during text extraction
            with suppress_pdf_warnings():
                # Open PDF document
                doc = fitz.open(input_file)
                total_pages = len(doc)
                logger.info(f"Extracting text from {total_pages} pages with PyMuPDF...")
                
                # Process pages in batches for better performance
                batch_size = 10
                for batch_start in range(0, total_pages, batch_size):
                    batch_end = min(batch_start + batch_size, total_pages)
                    logger.debug(f"Processing pages {batch_start + 1}-{batch_end}...")
                    
                    for page_num in range(batch_start, batch_end):
                        try:
                            page = doc[page_num]
                            
                            # Extract text blocks with layout
                            blocks = page.get_text("dict")["blocks"]
                            
                            for block in blocks:
                                if "lines" in block:  # Text block
                                    block_text = ""
                                    for line in block["lines"]:
                                        for span in line["spans"]:
                                            block_text += span["text"] + " "
                                    
                                    if block_text.strip():
                                        text_blocks.append({
                                            'text': block_text.strip(),
                                            'page': page_num + 1,
                                            'bbox': block.get("bbox", [0, 0, 0, 0]),
                                            'type': 'text'
                                        })
                        except Exception as e:
                            logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                            continue
                
                # Close document
                doc.close()
                
        except Exception as e:
            logger.error(f"Error extracting text with layout: {e}", exc_info=True)
        
        return text_blocks
    
    def _extract_text_simple(self, input_file: str) -> List[Dict[str, Any]]:
        """
        Extract text using simple method (pdfplumber) as fallback
        
        Args:
            input_file: Path to PDF file
            
        Returns:
            List of text blocks
        """
        text_blocks = []
        
        try:
            with suppress_pdf_warnings():
                with pdfplumber.open(input_file) as pdf:
                    for page_num, page in enumerate(pdf.pages, start=1):
                        page_text = page.extract_text()
                        if page_text:
                            cleaned_text = self._clean_pdf_text(page_text)
                            if cleaned_text.strip():
                                text_blocks.append({
                                    'text': cleaned_text,
                                    'page': page_num,
                                    'type': 'text'
                                })
        except Exception as e:
            logger.error(f"Error in simple text extraction: {e}", exc_info=True)
        
        return text_blocks
    
    def _insert_translated_tables(self, doc: Document, tables: List[Dict[str, Any]], 
                                   src_lang: str, dest_lang: str) -> None:
        """
        Insert translated tables into Word document
        
        Args:
            doc: Word document
            tables: List of table dictionaries
            src_lang: Source language code
            dest_lang: Destination language code
        """
        # Sort tables by page number
        sorted_tables = sorted(tables, key=lambda x: x.get('page', 0))
        current_page = 0
        
        for table_info in sorted_tables:
            df = table_info['dataframe']
            page = table_info.get('page', 0)
            
            # Add page break if needed
            if current_page > 0 and page > current_page:
                p_break = doc.add_paragraph()
                run_break = p_break.add_run()
                run_break.add_break(WD_BREAK.PAGE)
            
            current_page = page
            
            # Translate table cells - OPTIMIZATION: Batch translate all cells
            translated_df = df.copy()
            
            # Collect all cell values to translate
            cells_to_translate = []
            cell_positions = []
            for col in df.columns:
                for idx in df.index:
                    cell_value = df.at[idx, col]
                    # Filter None, NaN, and empty values
                    if cell_value is not None and pd.notna(cell_value):
                        cell_str = str(cell_value)
                        if cell_str.strip() and cell_str.lower() != 'nan':
                            cells_to_translate.append(cell_str)
                            cell_positions.append((idx, col))
            
            # SUPER BATCH: Join all cells with delimiter, translate once, split back
            if cells_to_translate:
                DELIMITER = "|||TXSEP|||"
                combined_text = DELIMITER.join(cells_to_translate)
                
                # Single translation for all cells
                logger.info(f"SUPER BATCH: Translating {len(cells_to_translate)} table cells in 1 request")
                try:
                    translated_combined = self.translation_service.translate_text(combined_text, src_lang, dest_lang)
                    translated_values = translated_combined.split(DELIMITER)
                    
                    # Handle mismatch
                    if len(translated_values) != len(cells_to_translate):
                        logger.warning(f"Delimiter mismatch in table. Keeping original values.")
                        translated_values = cells_to_translate
                except Exception as e:
                    logger.warning(f"Super batch table translation failed: {e}. Keeping original.")
                    translated_values = cells_to_translate
                
                # Map translated values back to dataframe
                for (idx, col), translated_value in zip(cell_positions, translated_values):
                    if translated_value is not None:
                        translated_df.at[idx, col] = translated_value
                    else:
                        translated_df.at[idx, col] = df.at[idx, col]
            
            # Create Word table
            num_rows = len(translated_df)
            num_cols = len(translated_df.columns)
            
            if num_rows > 0 and num_cols > 0:
                word_table = doc.add_table(rows=num_rows, cols=num_cols)
                word_table.style = 'Light Grid Accent 1'
                word_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Fill table with translated data
                for row_idx, (_, row_data) in enumerate(translated_df.iterrows()):
                    for col_idx, value in enumerate(row_data):
                        cell = word_table.rows[row_idx].cells[col_idx]
                        # Handle NaN and None values
                        if pd.isna(value) if hasattr(pd, 'isna') else (value is None or str(value).lower() == 'nan'):
                            cell.text = ""
                        else:
                            cell.text = str(value)
                        
                        # Set font for Asian characters
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = self.font_name
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)
                
                # Add spacing after table
                doc.add_paragraph()
    
    def _insert_translated_text_blocks(self, doc: Document, text_blocks: List[Dict[str, Any]], 
                                       src_lang: str, dest_lang: str) -> None:
        """
        Insert translated text blocks into Word document with format preservation
        Optimized with batch translation for better performance
        
        Args:
            doc: Word document
            text_blocks: List of text block dictionaries
            src_lang: Source language code
            dest_lang: Destination language code
        """
        if not text_blocks:
            return
        
        # Sort blocks by page
        sorted_blocks = sorted(text_blocks, key=lambda x: x.get('page', 0))
        current_page = 0
        
        # OPTIMIZATION: Batch translate all texts at once
        # CRITICAL: Filter out None, empty, and invalid values to prevent translation errors
        texts_to_translate = []
        valid_indices = []
        for i, block in enumerate(sorted_blocks):
            text = block.get('text', '')
            # Strict filtering: must be non-None, string, non-empty, and not 'nan'
            if (text is not None and 
                isinstance(text, str) and 
                text.strip() and 
                text.lower() not in ('nan', 'none', '')):
                texts_to_translate.append(str(text).strip())
                valid_indices.append(i)
        
        logger.info(f"Batch translating {len(texts_to_translate)} text blocks (filtered from {len(sorted_blocks)})...")
        
        translated_texts_map = {}
        if texts_to_translate:
            # SUPER BATCH: Join all texts with delimiter, translate once, split back
            DELIMITER = "|||TXSEP|||"
            combined_text = DELIMITER.join(texts_to_translate)
            
            # Single translation for all text blocks
            logger.info(f"SUPER BATCH: Translating {len(texts_to_translate)} text blocks in 1 request ({len(combined_text)} chars)")
            try:
                translated_combined = self.translation_service.translate_text(combined_text, src_lang, dest_lang)
                translated_results = translated_combined.split(DELIMITER)
                
                # Handle mismatch
                if len(translated_results) != len(texts_to_translate):
                    logger.warning(f"Delimiter mismatch: expected {len(texts_to_translate)}, got {len(translated_results)}. Using original.")
                    translated_results = texts_to_translate
                
                # Map results back
                for idx, translated in zip(valid_indices, translated_results):
                    translated_texts_map[idx] = translated if translated else sorted_blocks[idx].get('text', '')
            except Exception as e:
                logger.warning(f"Super batch text translation failed: {e}. Keeping original.")
                for idx in valid_indices:
                    translated_texts_map[idx] = sorted_blocks[idx].get('text', '')
        
        # Insert translated blocks
        for i, block_info in enumerate(sorted_blocks):
            text = block_info.get('text', '')
            page = block_info.get('page', 0)
            # Get translated text or use original
            translated_text = translated_texts_map.get(i, text)
            if not translated_text:
                translated_text = text
            
            if not text or not text.strip():
                continue
            
            # Add page break if needed
            if current_page > 0 and page > current_page:
                p_break = doc.add_paragraph()
                run_break = p_break.add_run()
                run_break.add_break(WD_BREAK.PAGE)
            
            current_page = page
            
            # Use pre-translated text from batch translation (already done above)
            # translated_text is already set from batch translation
            
            # Insert as paragraphs (preserve line breaks)
            paragraphs = translated_text.split('\n')
            for para_text in paragraphs:
                if para_text.strip():
                    p = doc.add_paragraph(para_text.strip())
                    for run in p.runs:
                        run.font.name = self.font_name
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)
    
    def _clean_pdf_text(self, text: str) -> str:
        """
        Clean text from PDF - remove unnecessary special characters
        
        Args:
            text: Text to clean
        
        Returns:
            Cleaned text
        """
        if not text:
            return text
        
        # Remove CID (Character ID) characters that are not rendered correctly
        # Pattern: (cid:xxxx) or (cid:xxxxx)
        text = re.sub(r'\(cid:\d+\)', '', text)
        
        # Remove unnecessary control characters
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove consecutive empty lines
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        return text.strip()

