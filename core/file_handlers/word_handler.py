"""
Word file handler for translation with complete coverage and QA reporting.
"""
import re
import os
from typing import Any, Dict, List, Optional, Set
from docx import Document
from docx.text.paragraph import Paragraph

from translation_app.config import config
from translation_app.core.file_translation_control import FileTranslationInterrupted, FileTranslationStopRequested
from translation_app.core.ocr_handler import get_ocr_handler
from translation_app.core.translator import TranslationService
from translation_app.utils.error_handler import FileProcessingError
from translation_app.utils.logger import logger
from translation_app.core.translation_job import redact_sensitive


class WordTranslationTarget:
    """Represents a text segment or paragraph to be translated in a Word document."""

    def __init__(
        self,
        target_id: str,
        location: str,
        text: str,
        paragraph: Paragraph,
        can_translate: bool = True,
        skip_reason: Optional[str] = None,
    ):
        self.target_id = target_id
        self.location = location  # 'body', 'table', 'nested_table', 'header', 'footer', 'textbox', 'unsupported'
        self.text = text
        self.paragraph = paragraph
        self.can_translate = can_translate
        self.skip_reason = skip_reason
        self.formatting_risk = "low"


class WordHandler:
    """Handler for Word file translation with completeness hardening."""

    _URL_RE = re.compile(r"^(?:https?://|www\.)\S+$", re.IGNORECASE)
    _EMAIL_RE = re.compile(r"^[^@\s]+@[^@\s]+\.[^@\s]+$")
    # Strict Unix path requiring at least one sub-folder (e.g. /home/user), Windows path, or UNC path
    _FILE_PATH_RE = re.compile(r"^(?:[A-Za-z]:\\|\\\\)(?:[^\\]+\\)*[^\\]+$|^(?:/[A-Za-z0-9_-]+){2,}(?:\.[A-Za-z0-9]+)?$")
    _FIELD_HINT_RE = re.compile(r"\b(?:HYPERLINK|MERGEFIELD|PAGE|NUMPAGES|TOC)\b", re.IGNORECASE)
    _FLIGHT_CODE_RE = re.compile(r"^[A-Z]{1,4}\d{1,4}[A-Z]?$")
    _NUMERICISH_TOKEN_RE = re.compile(r"^(?=.*\d)[\d\s.,:/()%+\-]+$")
    _FRAGMENTED_RUN_MIN_COUNT = 2
    _FRAGMENTED_RUN_MIN_TOTAL_CHARS = 20

    def __init__(self, translation_service: TranslationService):
        self.translation_service = translation_service
        self.ocr_handler = get_ocr_handler()
        self.progress_callback = None
        self.last_word_qa_report: Optional[Dict[str, Any]] = None

    def translate(self, input_file: str, output_file: str, src_lang: str, dest_lang: str) -> Dict[str, Any]:
        """Translate a DOCX file while preserving all core formatting and layout structure."""
        doc = None
        try:
            logger.info(f"Starting Word translation: {input_file}")
            if self.progress_callback:
                self.progress_callback("Reading Word document...", 10)

            file_size_mb = self._get_file_size_mb(input_file)
            if file_size_mb > config.warning_file_size_mb:
                logger.warning(f"Large file detected: {file_size_mb:.1f}MB")

            doc = Document(input_file)

            # 1. Collect all translation targets systematically
            if self.progress_callback:
                self.progress_callback("Analyzing document structure...", 20)
            targets = self.collect_word_translation_targets(doc)

            # 2. Setup QA Reporting metrics
            report = {
                "total_candidates": len(targets),
                "translated_candidates": 0,
                "skipped_candidates": 0,
                "failed_candidates": 0,
                "by_location": {
                    "body": 0,
                    "table": 0,
                    "nested_table": 0,
                    "header": 0,
                    "footer": 0,
                    "textbox": 0,
                    "unsupported": 0,
                },
                "skip_reasons": {},
                "unsupported_locations": [],
            }

            # 3. Translate candidates
            if self.progress_callback:
                self.progress_callback("Translating document content...", 50)

            api_requests = 0
            for idx, target in enumerate(targets):
                getattr(self.translation_service, "raise_if_file_translation_stopped", lambda: None)()
                # Update location statistics
                loc = target.location
                report["by_location"][loc] = report["by_location"].get(loc, 0) + 1

                if not target.can_translate:
                    report["skipped_candidates"] += 1
                    reason = target.skip_reason or "unknown"
                    report["skip_reasons"][reason] = report["skip_reasons"].get(reason, 0) + 1
                    continue

                # Perform the actual translation on the target runs
                success = self._translate_target_runs(target, src_lang, dest_lang)
                if success:
                    report["translated_candidates"] += 1
                    api_requests += 1
                else:
                    report["failed_candidates"] += 1

                # Periodically update progress callback
                if self.progress_callback and len(targets) > 0:
                    percent = int(50 + (idx / len(targets)) * 40)
                    self.progress_callback(f"Translating text segments ({idx+1}/{len(targets)})...", percent)

            # Save the clean QA report safely
            self.last_word_qa_report = redact_sensitive(report)

            # Strict error checking policy
            failed_cand = report.get("failed_candidates", 0)
            trans_cand = report.get("translated_candidates", 0)
            total_cand = report.get("total_candidates", 0)

            if failed_cand > 0:
                if trans_cand == 0:
                    raise FileProcessingError(
                        "Không dịch được nội dung Word nào. Vui lòng kiểm tra provider/API key/chế độ dịch."
                    )
                else:
                    raise FileProcessingError(
                        f"Phát hiện lỗi dịch thuật trên {failed_cand}/{total_cand} đoạn văn bản. Vui lòng kiểm tra API key hoặc kết nối mạng."
                    )

            if self.progress_callback:
                self.progress_callback("Saving translated Word document...", 95)
            doc.save(output_file)

            logger.info(f"Word translation completed: {output_file}")
            if self.progress_callback:
                self.progress_callback("Done", 100)

            return {
                "api_requests": api_requests,
                "images_processed": 0,
                "images_skipped": 0,
            }
        except FileTranslationStopRequested as exc:
            partial_saved, save_error = self._save_partial_document(doc, output_file, exc.status)
            raise FileTranslationInterrupted(
                exc.status,
                output_file=output_file,
                partial_saved=partial_saved,
                save_error=save_error,
            ) from exc
        except FileProcessingError as exc:
            logger.error(str(exc))
            raise exc
        except Exception as exc:
            error_msg = f"Error translating Word file: {exc}"
            logger.error(error_msg)
            raise FileProcessingError(error_msg, original_error=exc) from exc

    def collect_word_translation_targets(self, doc: Document) -> List[WordTranslationTarget]:
        """Systematically collect all paragraphs and text segments to be translated."""
        processed_element_ids: Set[int] = set()
        targets: List[WordTranslationTarget] = []

        def _add_paragraph_target(p: Paragraph, location: str):
            if id(p._element) in processed_element_ids:
                return
            processed_element_ids.add(id(p._element))

            text = p.text
            stripped = (text or "").strip()

            # Classify skipping status and skip reason
            can_translate = True
            skip_reason = None

            if not stripped:
                can_translate = False
                skip_reason = "empty_or_whitespace"
            elif self._should_skip_text(stripped):
                can_translate = False
                skip_reason = "matches_skip_patterns"

            target_id = f"word_{location}_{len(targets)}"
            targets.append(
                WordTranslationTarget(
                    target_id=target_id,
                    location=location,
                    text=text,
                    paragraph=p,
                    can_translate=can_translate,
                    skip_reason=skip_reason,
                )
            )

        # 1. Collect body paragraphs
        for p in doc.paragraphs:
            _add_paragraph_target(p, "body")

        # 2. Collect tables recursively
        def _traverse_table(table, location: str = "table"):
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        _add_paragraph_target(p, location)
                    for nested in cell.tables:
                        _traverse_table(nested, "nested_table")

        for table in doc.tables:
            _traverse_table(table)

        # 3. Collect headers
        for section in doc.sections:
            if section.header is not None:
                for p in section.header.paragraphs:
                    _add_paragraph_target(p, "header")

        # 4. Collect footers
        for section in doc.sections:
            if section.footer is not None:
                for p in section.footer.paragraphs:
                    _add_paragraph_target(p, "footer")

        # 5. Collect textboxes via XML parsing using document namespace maps
        try:
            from lxml import etree
            namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            xpath_txbx = etree.XPath("//w:txbxContent", namespaces=namespaces)
            xpath_p = etree.XPath(".//w:p", namespaces=namespaces)
            txbx_contents = xpath_txbx(doc.element)
            for txbx in txbx_contents:
                for p_elem in xpath_p(txbx):
                    if id(p_elem) not in processed_element_ids:
                        p_obj = Paragraph(p_elem, doc)
                        _add_paragraph_target(p_obj, "textbox")
        except Exception as exc:
            logger.warning(f"Failed to scan XML textboxes: {exc}")

        return targets

    def _translate_target_runs(self, target: WordTranslationTarget, src_lang: str, dest_lang: str) -> bool:
        """Translate individual runs of a target paragraph while preserving run-level formatting."""
        getattr(self.translation_service, "raise_if_file_translation_stopped", lambda: None)()
        original_text = target.text
        if self._should_skip_text(original_text):
            return False

        # If the paragraph has no runs but has text, translate paragraph text directly
        if not target.paragraph.runs and original_text.strip():
            try:
                translated_text = self.translation_service.translate_long_text(
                    original_text,
                    src_lang,
                    dest_lang,
                )
                if translated_text != original_text:
                    target.paragraph.text = translated_text
                    return True
                return self._is_unchanged_translation_acceptable(original_text)
            except Exception as exc:
                logger.warning(f"Error translating Word paragraph: {exc}")
                return False

        eligible_runs = self._collect_translatable_runs(target.paragraph.runs)
        if self._should_translate_fragmented_runs_as_block(target.paragraph.runs, eligible_runs):
            return self._translate_fragmented_paragraph(target, src_lang, dest_lang)

        translated_any = False
        acceptable_unchanged = False
        all_failed = True
        has_run = False

        for run in eligible_runs:
            getattr(self.translation_service, "raise_if_file_translation_stopped", lambda: None)()
            run_text = run.text

            has_run = True
            try:
                translated_text = self.translation_service.translate_long_text(
                    run_text,
                    src_lang,
                    dest_lang,
                )
                if translated_text != run_text:
                    run.text = translated_text
                    translated_any = True
                elif self._is_unchanged_translation_acceptable(run_text):
                    acceptable_unchanged = True
                all_failed = False
            except Exception as exc:
                logger.warning(f"Error translating Word run in target {target.target_id}: {exc}")

        if not has_run:
            return False

        # If there were runs to translate and they all failed, report translation failure for this segment
        if all_failed:
            return False

        return translated_any or acceptable_unchanged or not has_run

    def _collect_translatable_runs(self, runs: List[Any]) -> List[Any]:
        eligible_runs: List[Any] = []
        for run in runs:
            run_text = run.text
            if self._should_skip_text(run_text):
                continue
            if self._run_has_field_markup(run):
                continue
            eligible_runs.append(run)
        return eligible_runs

    def _should_translate_fragmented_runs_as_block(self, all_runs: List[Any], eligible_runs: List[Any]) -> bool:
        if len(eligible_runs) < self._FRAGMENTED_RUN_MIN_COUNT:
            return False

        for run in all_runs:
            run_text = run.text or ""
            if not run_text.strip():
                continue
            if self._run_has_field_markup(run):
                return False
            if self._should_skip_text(run_text):
                return False

        total_chars = sum(len((run.text or "").strip()) for run in eligible_runs)
        return total_chars >= self._FRAGMENTED_RUN_MIN_TOTAL_CHARS

    def _translate_fragmented_paragraph(self, target: WordTranslationTarget, src_lang: str, dest_lang: str) -> bool:
        getattr(self.translation_service, "raise_if_file_translation_stopped", lambda: None)()
        original_text = target.text
        try:
            translated_text = self.translation_service.translate_long_text(
                original_text,
                src_lang,
                dest_lang,
            )
        except Exception as exc:
            logger.warning(f"Error translating fragmented Word paragraph in target {target.target_id}: {exc}")
            return False

        if translated_text == original_text:
            return self._is_unchanged_translation_acceptable(original_text)

        first_run = None
        for run in target.paragraph.runs:
            if self._run_has_field_markup(run):
                continue
            if first_run is None:
                first_run = run
                run.text = translated_text
                continue
            run.text = ""
        if first_run is None:
            target.paragraph.text = translated_text
        return True

    def _save_partial_document(self, doc: Optional[Document], output_file: str, status: str) -> tuple[bool, Optional[Exception]]:
        if doc is None:
            return False, None
        try:
            if self.progress_callback:
                self.progress_callback(f"Saving partial Word document ({status})...", 95)
            doc.save(output_file)
            logger.info(f"Saved partial Word output after file translation was {status}: {output_file}")
            return True, None
        except Exception as exc:
            logger.error(f"Failed to save partial Word output after {status}: {exc}")
            return False, exc

    def _should_skip_text(self, text: str) -> bool:
        """Robust skip logic to ensure symbols, paths, fields, and URLs are skipped, but normal text is preserved."""
        stripped = (text or "").strip()
        if not stripped:
            return True

        # Skip punctuation-only, divider, or pure symbol strings that lack any alphanumeric character
        if not any(c.isalnum() for c in stripped):
            return True

        # Match strict URLs
        if self._URL_RE.match(stripped):
            return True

        # Match emails
        if self._EMAIL_RE.match(stripped):
            return True

        # Match strict local or UNC file paths
        if self._FILE_PATH_RE.match(stripped):
            return True

        # Match Word Field Hints
        if self._FIELD_HINT_RE.search(stripped):
            return True

        if self._is_nonlinguistic_ascii_token(stripped):
            return True

        return False

    def _is_nonlinguistic_ascii_token(self, stripped: str) -> bool:
        if self._FLIGHT_CODE_RE.match(stripped):
            return True

        if self._NUMERICISH_TOKEN_RE.match(stripped):
            return True

        if not stripped.isascii() or not any(ch.isdigit() for ch in stripped):
            return False

        words = re.findall(r"[A-Za-z]+", stripped)
        return bool(words) and all(len(word) <= 2 for word in words)

    def _is_unchanged_translation_acceptable(self, text: str) -> bool:
        stripped = (text or "").strip()
        if not stripped:
            return True
        if self._is_nonlinguistic_ascii_token(stripped):
            return True
        return bool(re.fullmatch(r"[A-Za-z]", stripped))

    def _run_has_field_markup(self, run) -> bool:
        """Check if the run XML contains Word fields or instructional text markup."""
        try:
            xml = run._element.xml
        except Exception:
            return False
        return "w:fldChar" in xml or "w:instrText" in xml

    def _get_file_size_mb(self, file_path: str) -> float:
        return os.path.getsize(file_path) / (1024 * 1024)
