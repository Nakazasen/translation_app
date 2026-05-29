import concurrent.futures
from pathlib import Path

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from translation_app.core.file_translation_control import FileTranslationInterrupted, FileTranslationStopRequested
from translation_app.core.file_handlers.word_handler import WordHandler
from tests.test_word_ppt_format_preservation import FakeTranslationService
from translation_app.utils.error_handler import FileProcessingError


def _add_textbox_to_doc(doc, text: str):
    """Add a raw w:txbxContent textbox element to document body XML for testing."""
    p = doc.add_paragraph()
    # Create an inline drawing shape wrapper or a raw textbox block
    # We can insert a raw w:txbxContent element directly into paragraph or body
    # Let's insert a w:txbxContent inside a paragraph's XML for simple test traversal
    txbx = OxmlElement('w:txbxContent')
    txbx_p = OxmlElement('w:p')
    txbx_r = OxmlElement('w:r')
    txbx_t = OxmlElement('w:t')
    txbx_t.text = text
    txbx_r.append(txbx_t)
    txbx_p.append(txbx_r)
    txbx.append(txbx_p)
    p._element.append(txbx)


def test_word_translates_all_body_paragraphs(tmp_path):
    input_path = tmp_path / "body.docx"
    output_path = tmp_path / "body_out.docx"
    service = FakeTranslationService()

    doc = Document()
    doc.add_paragraph("Đoạn văn thứ nhất.")
    doc.add_paragraph("Đoạn văn thứ hai.")
    doc.add_paragraph("Đoạn văn thứ ba.")
    doc.save(input_path)

    handler = WordHandler(service)
    try:
        stats = handler.translate(str(input_path), str(output_path), "vi", "ja")
    finally:
        service.executor.shutdown(wait=True)

    assert Path(output_path).exists()
    result = Document(output_path)
    assert len(result.paragraphs) == 3
    assert "Đoạn văn thứ nhất.-ja" in result.paragraphs[0].text
    assert "Đoạn văn thứ hai.-ja" in result.paragraphs[1].text
    assert "Đoạn văn thứ ba.-ja" in result.paragraphs[2].text

    report = handler.last_word_qa_report
    assert report is not None
    assert report["total_candidates"] >= 3
    assert report["translated_candidates"] == 3
    assert report["skipped_candidates"] >= 0
    assert report["failed_candidates"] == 0
    assert report["by_location"]["body"] == 3


def test_word_translates_table_and_nested_table_paragraphs(tmp_path):
    input_path = tmp_path / "tables.docx"
    output_path = tmp_path / "tables_out.docx"
    service = FakeTranslationService()

    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = "Nội dung bảng chính"
    
    # Add nested table inside the cell
    nested = cell.add_table(rows=1, cols=1)
    nested.cell(0, 0).text = "Nội dung bảng con"
    doc.save(input_path)

    handler = WordHandler(service)
    try:
        stats = handler.translate(str(input_path), str(output_path), "vi", "ja")
    finally:
        service.executor.shutdown(wait=True)

    result = Document(output_path)
    assert "Nội dung bảng chính-ja" in result.tables[0].cell(0, 0).paragraphs[0].text
    # Note: cell.paragraphs[0] might contain the text, nested table is in cell.tables[0]
    assert "Nội dung bảng con-ja" in result.tables[0].cell(0, 0).tables[0].cell(0, 0).text

    report = handler.last_word_qa_report
    assert report["by_location"]["table"] >= 1
    assert report["by_location"]["nested_table"] >= 1


def test_word_translates_headers_and_footers(tmp_path):
    input_path = tmp_path / "header_footer.docx"
    output_path = tmp_path / "header_footer_out.docx"
    service = FakeTranslationService()

    doc = Document()
    section = doc.sections[0]
    header = section.header
    header.paragraphs[0].text = "Tiêu đề đầu trang"
    footer = section.footer
    footer.paragraphs[0].text = "Tiêu đề cuối trang"
    doc.save(input_path)

    handler = WordHandler(service)
    try:
        stats = handler.translate(str(input_path), str(output_path), "vi", "ja")
    finally:
        service.executor.shutdown(wait=True)

    result = Document(output_path)
    assert "Tiêu đề đầu trang-ja" in result.sections[0].header.paragraphs[0].text
    assert "Tiêu đề cuối trang-ja" in result.sections[0].footer.paragraphs[0].text

    report = handler.last_word_qa_report
    assert report["by_location"]["header"] >= 1
    assert report["by_location"]["footer"] >= 1


def test_word_does_not_silently_skip_textboxes(tmp_path):
    input_path = tmp_path / "textbox.docx"
    output_path = tmp_path / "textbox_out.docx"
    service = FakeTranslationService()

    doc = Document()
    _add_textbox_to_doc(doc, "Nội dung textbox")
    doc.save(input_path)

    handler = WordHandler(service)
    try:
        stats = handler.translate(str(input_path), str(output_path), "vi", "ja")
    finally:
        service.executor.shutdown(wait=True)

    report = handler.last_word_qa_report
    assert report["by_location"]["textbox"] >= 1
    # Check that either it got translated or reported in QA skipped list (never silently ignored)
    assert report["total_candidates"] >= 2  # 1 body paragraph container, 1 textbox paragraph


def test_word_skip_logic_does_not_skip_normal_vietnamese():
    service = FakeTranslationService()
    handler = WordHandler(service)

    # Valid Vietnamese strings that should NOT be skipped
    assert not handler._should_skip_text("Dự án Dịch Thuật")
    assert not handler._should_skip_text("Chào thế giới!")
    assert not handler._should_skip_text("1. Hướng dẫn sử dụng")
    assert not handler._should_skip_text("/ Dịch thuật")  # forward slash prefix shouldn't skip valid text
    assert not handler._should_skip_text("Tiếng Việt có dấu á à ả ã ạ")

    # Strings that SHOULD be skipped
    assert handler._should_skip_text("https://google.com")
    assert handler._should_skip_text("contact@example.com")
    assert handler._should_skip_text("C:\\temp\\file.txt")
    assert handler._should_skip_text("•")  # non-alphanumeric single symbol
    assert handler._should_skip_text("---")  # divider
    assert handler._should_skip_text("VJ767")
    assert handler._should_skip_text("7kg + 20kg")
    assert handler._should_skip_text("0.000")
    assert handler._should_skip_text("/06")
    assert not handler._should_skip_text("Day 1")


def test_word_ai_all_candidates_fail_raises_not_success(tmp_path):
    input_path = tmp_path / "error.docx"
    output_path = tmp_path / "error_out.docx"

    class FailingTranslationService(FakeTranslationService):
        def translate_long_text(self, text, src_lang, dest_lang):
            raise Exception("API Limit exceeded")

    service = FailingTranslationService()
    doc = Document()
    doc.add_paragraph("Đoạn văn lỗi.")
    doc.save(input_path)

    handler = WordHandler(service)
    try:
        with pytest.raises(FileProcessingError) as exc_info:
            handler.translate(str(input_path), str(output_path), "vi", "ja")
        
        assert "Không dịch được nội dung Word nào" in str(exc_info.value)
    finally:
        service.executor.shutdown(wait=True)

    report = handler.last_word_qa_report
    assert report is not None
    assert report["failed_candidates"] == 1
    assert report["translated_candidates"] == 0


def test_word_partial_failure_is_not_silent(tmp_path):
    input_path = tmp_path / "partial.docx"
    output_path = tmp_path / "partial_out.docx"

    class PartialFailingTranslationService(FakeTranslationService):
        def __init__(self):
            super().__init__()
            self.calls = 0

        def translate_long_text(self, text, src_lang, dest_lang):
            self.calls += 1
            if self.calls == 1:
                return text + "-ja"
            raise Exception("API quota exceeded")

    service = PartialFailingTranslationService()
    doc = Document()
    doc.add_paragraph("Đoạn văn dịch được.")
    doc.add_paragraph("Đoạn văn lỗi.")
    doc.save(input_path)

    handler = WordHandler(service)
    try:
        with pytest.raises(FileProcessingError) as exc_info:
            handler.translate(str(input_path), str(output_path), "vi", "ja")
        
        assert "Phát hiện lỗi dịch thuật trên 1" in str(exc_info.value)
    finally:
        service.executor.shutdown(wait=True)

    report = handler.last_word_qa_report
    assert report is not None
    assert report["failed_candidates"] == 1
    assert report["translated_candidates"] == 1


def test_word_pause_saves_partial_output(tmp_path):
    input_path = tmp_path / "pause.docx"
    output_path = tmp_path / "pause_out.docx"

    class PausingTranslationService(FakeTranslationService):
        def __init__(self):
            super().__init__()
            self.stop_checks = 0

        def raise_if_file_translation_stopped(self):
            self.stop_checks += 1
            if self.stop_checks >= 4:
                raise FileTranslationStopRequested("paused")

    service = PausingTranslationService()
    doc = Document()
    doc.add_paragraph("Doan 1")
    doc.add_paragraph("Doan 2")
    doc.save(input_path)

    handler = WordHandler(service)
    try:
        with pytest.raises(FileTranslationInterrupted) as exc_info:
            handler.translate(str(input_path), str(output_path), "vi", "ja")
    finally:
        service.executor.shutdown(wait=True)

    assert exc_info.value.status == "paused"
    assert exc_info.value.partial_saved is True
    assert Path(output_path).exists()

    result = Document(output_path)
    assert result.paragraphs[0].text == "Doan 1-ja"
    assert result.paragraphs[1].text == "Doan 2"


def test_word_google_success_path_still_passes(tmp_path):
    input_path = tmp_path / "google_success.docx"
    output_path = tmp_path / "google_success_out.docx"
    service = FakeTranslationService()

    doc = Document()
    doc.add_paragraph("Đoạn văn Google dịch 1.")
    doc.add_paragraph("Đoạn văn Google dịch 2.")
    doc.save(input_path)

    handler = WordHandler(service)
    try:
        stats = handler.translate(str(input_path), str(output_path), "vi", "ja")
        assert stats["api_requests"] == 2
    finally:
        service.executor.shutdown(wait=True)

    report = handler.last_word_qa_report
    assert report is not None
    assert report["failed_candidates"] == 0
    assert report["translated_candidates"] == 2


def test_word_nonlinguistic_tokens_are_skipped_not_failed(tmp_path):
    input_path = tmp_path / "codes.docx"
    output_path = tmp_path / "codes_out.docx"
    service = FakeTranslationService()

    doc = Document()
    doc.add_paragraph("VJ767")
    doc.add_paragraph("7kg + 20kg")
    doc.add_paragraph("0.000")
    doc.add_paragraph("/06")
    doc.save(input_path)

    handler = WordHandler(service)
    try:
        handler.translate(str(input_path), str(output_path), "vi", "ja")
    finally:
        service.executor.shutdown(wait=True)

    result = Document(output_path)
    assert [paragraph.text for paragraph in result.paragraphs] == ["VJ767", "7kg + 20kg", "0.000", "/06"]
    assert service.calls == []

    report = handler.last_word_qa_report
    assert report is not None
    assert report["failed_candidates"] == 0
    assert report["translated_candidates"] == 0
    assert report["skipped_candidates"] >= 4


def test_word_acceptable_unchanged_tokens_do_not_fail_qa(tmp_path):
    input_path = tmp_path / "markers.docx"
    output_path = tmp_path / "markers_out.docx"

    class UnchangedTranslationService(FakeTranslationService):
        def translate_long_text(self, text, src_lang, dest_lang):
            self.calls.append(text)
            return text

    service = UnchangedTranslationService()
    doc = Document()
    doc.add_paragraph("H")
    doc.add_paragraph("l")
    doc.save(input_path)

    handler = WordHandler(service)
    try:
        handler.translate(str(input_path), str(output_path), "vi", "ja")
    finally:
        service.executor.shutdown(wait=True)

    result = Document(output_path)
    assert [paragraph.text for paragraph in result.paragraphs] == ["H", "l"]

    report = handler.last_word_qa_report
    assert report is not None
    assert report["failed_candidates"] == 0
    assert report["translated_candidates"] == 2


def test_word_qa_report_public_safe(tmp_path):
    input_path = tmp_path / "safe.docx"
    output_path = tmp_path / "safe_out.docx"
    service = FakeTranslationService()

    doc = Document()
    doc.add_paragraph("AIzaSy-API-KEY-SECRET")
    doc.save(input_path)

    handler = WordHandler(service)
    try:
        stats = handler.translate(str(input_path), str(output_path), "vi", "ja")
    finally:
        service.executor.shutdown(wait=True)

    report = handler.last_word_qa_report
    report_str = repr(report)

    # QA report must be completely clean of keys, prompts, and raw text
    assert "AIzaSy" not in report_str
    assert "SECRET" not in report_str
    assert "api_key" not in report_str
